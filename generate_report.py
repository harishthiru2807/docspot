from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # Title Section
    title = doc.add_heading('DocSpot', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('AI-Powered Universal Telemedicine Ecosystem')
    run = subtitle.runs[0]
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(37, 99, 235) # Trust Blue
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('Integrated Healthcare Portal Management System').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Abstract
    doc.add_heading('1. Abstract', level=1)
    doc.add_paragraph(
        "DocSpot is a decentralized, AI-driven healthcare platform designed to bridge the rural-urban medical divide. "
        "By integrating a bilingual (Tamil & English) AI Video Doctor, role-based clinical dashboards, and real-time inventory synchronization, "
        "DocSpot provides a seamless healthcare journey from symptom onset to medication fulfillment. "
        "The platform leverages modern web technologies to ensure high performance, security, and accessibility, "
        "making professional medical advice available to every citizen, regardless of their location or linguistic background."
    )

    # Introduction
    doc.add_heading('2. Introduction', level=1)
    doc.add_paragraph(
        "Healthcare accessibility in India is hampered by a significant shortage of specialists in rural areas and the presence "
        "of profound linguistic barriers in clinical communication. Patients in remote villages often struggle to describe "
        "symptoms accurately or lack a reliable means to maintain medical history."
    )

    # Objectives
    doc.add_heading('3. Objectives', level=1)
    objectives = [
        ("Accessibility", "Provide 24/7 AI-driven medical guidance in regional languages (Tamil)."),
        ("Stakeholder Integration", "Synchronize clinical data across patients, doctors, nurses, and pharmacies."),
        ("Trust & Interaction", "Humanize AI interaction through a realistic digital human avatar (Dr. ARIA)."),
        ("Data Digitization", "Move away from paper-based records to secure, QR-based digital health IDs."),
        ("Efficiency", "Reduce diagnostic latency through automated symptom checking and real-time ward monitoring.")
    ]
    for obj, desc in objectives:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{obj}: ")
        run.bold = True
        p.add_run(desc)

    # Technical Stack
    doc.add_heading('4. Technical Stack', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technology'
    hdr_cells[2].text = 'Reasoning'

    tech_stack = [
        ('Frontend', 'React.js & Vite', 'Component-based modularity and lightning-fast rendering.'),
        ('Backend', 'Node.js & Express', 'Non-blocking I/O for high-concurrency events.'),
        ('Security', 'JWT', 'Stateless authentication perfect for rural deployment.'),
        ('API Layer', 'Axios', 'Custom interceptors for seamless token management.')
    ]
    for layer, tech, reason in tech_stack:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = reason

    # Portals
    doc.add_heading('5. Functional Modules', level=1)
    portals = [
        ("Patient Portal", "Enables AI-driven diagnosis via Dr. ARIA, QR profiles, and vitals tracking."),
        ("Doctor Portal", "Queue management system with live vitals and video consultation."),
        ("Nurse Portal", "Ward oversight with real-time clinical alerts."),
        ("Pharmacy Portal", "Inventory synchronization with low-stock replenishment alerts.")
    ]
    for name, desc in portals:
        doc.add_heading(name, level=2)
        doc.add_paragraph(desc)

    # Innovation
    doc.add_heading('6. Innovation: AI Video Doctor (Dr. ARIA)', level=1)
    doc.add_paragraph(
        "The core innovation of DocSpot is the Digital Human Interface. Unlike traditional text bots, Dr. ARIA uses a "
        "custom CSS lip-sync engine coordinated with the Web Speech API. This high-fidelity interaction builds clinical "
        "trust, especially important for users in rural areas who may be skeptical of non-human interactions."
    )

    # Conclusion
    doc.add_heading('7. Conclusion', level=1)
    doc.add_paragraph(
        "DocSpot effectively demonstrates how AI and modern web architecture can solve real-world healthcare accessibility "
        "problems. By centralizing the workflow into a single ecosystem, we eliminate fragmented communication and provide "
        "a futuristic, scalable model for national healthcare delivery."
    )

    doc.save(r'C:\Users\haris\.gemini\antigravity\brain\37d27107-5f08-4408-b200-d7ba666918e4\DocSpot_Project_Report.docx')
    print("Report generated successfully.")

if __name__ == "__main__":
    create_report()
